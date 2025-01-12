import os
import PyPDF2
import pdfplumber
from docx import Document
import json
import psycopg2
import chromadb
import shutil
from dotenv import load_dotenv
import logging
from sentence_transformers import SentenceTransformer
import ollama
load_dotenv()


db_name=os.getenv("POSTGRES_DB")
db_user=os.getenv("POSTGRES_USER")
db_port=os.getenv("POSTGRES_PORT")
db_host=os.getenv("HOST")
db_password=os.getenv("POSTGRES_PASSWORD")
ollama_model=os.getenv("MODEL")

sentenct_transformer = os.getenv("TRANSFORMER1")
embedding_model = SentenceTransformer(sentenct_transformer)
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("processing.log"),
        logging.StreamHandler()
    ]
)

def connect_to_postgres():
    return psycopg2.connect(
        database=db_name,
        user=db_user,
        password=db_password,
        host=db_host,
        port=db_port
    )
    
def check_file_type(file_path):
    """
    Check the type of the file uploaded by the user
    """
    _,extension=os.path.splitext(file_path)

    return extension

def extract_content(file_path):
    """
    Check type and then extract content
    """
    file_extension=check_file_type(file_path)
    # print(file_extension)
    
    try:
        if file_extension == ".pdf":
            return extract_text_from_pdf(file_path)
        elif file_extension == ".docx":
            return extract_text_from_docx(file_path)
        elif file_extension == ".txt":
            return extract_text_from_txt(file_path)
        else:
            raise ValueError("Unsupported file type.Excepted file type PDF, DOCX, or TXT file.")
    except Exception as e:
        return f"Error processing file: {e}"
    
def extract_text_from_pdf(file_path):
    
    text=""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    
    doc = Document(file_path)
    print("DOC Is printint  ------- ",doc)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    # print("TEXT ******* ",text[:50])
    return text

def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        text = file.read()
    return text


# processing_status = {}

def save_uploaded_files(files, upload_dir):
    file_paths = []
    for file_path in files:
        filename = os.path.basename(file_path)
        destination_path = os.path.join(upload_dir, filename)
        shutil.move(file_path, destination_path)
        file_paths.append(destination_path)
        logging.info(f"FILE MOVED: {file_path} -> {destination_path}")  
    return file_paths

def insert_candidate_into_postgres(candidate_data):
    conn = connect_to_postgres()
    cursor = conn.cursor()

    # Insert candidate data
    insert_query = """
    INSERT INTO candidates (
        domain, name, email, phone, linkedin, github, professional_summary,
        technical_skills, work_experience, education, certifications, projects
    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    RETURNING candidate_id;
    """

    cursor.execute(insert_query, (
        candidate_data.get("domain", "Unknown"),
        candidate_data.get("name", ""),
        candidate_data.get("email", ""),
        candidate_data.get("phone", ""),
        candidate_data.get("linkedin", ""),
        candidate_data.get("github", ""),
        candidate_data.get("professional_summary", ""),
        candidate_data.get("technical_skills", []),
        json.dumps(candidate_data.get("work_experience", [])),
        json.dumps(candidate_data.get("education", [])),
        candidate_data.get("certifications", []),
        json.dumps(candidate_data.get("projects", []))
    ))
    candidate_id = cursor.fetchone()[0]
    conn.commit()
    cursor.close()
    conn.close()
    return candidate_id

def prepare_embedding_text(candidate):
    work_experience = " ".join(
        f"Role: {job.get('role', 'N/A')}, Company: {job.get('company', 'N/A')}, Responsibilities: {', '.join(job.get('responsibilities', []))}"
        for job in candidate.get('work_experience', [])
    )
    projects = " ".join(
        f"Project: {project.get('name', 'N/A')}, Description: {project.get('description', 'N/A')}"
        for project in candidate.get('projects', [])
    )
    professional_summary = candidate.get('professional_summary', 'No professional summary provided.')
    return f"{professional_summary} {work_experience} {projects}"

# Initialize ChromaDB
DB="candidate_resumes"
client = chromadb.PersistentClient(path=DB)
collection_name = "candidates_data"
existing_collection_names = [collection.name for collection in client.list_collections()]

if collection_name not in existing_collection_names:
    collection = client.create_collection(collection_name)
else:collection=client.get_collection(collection_name)

system_prompt="""You are an assistant that extracts structured information from resumes into JSON format just give me final json no other text. In addition to extracting details like name, email, and phone number, infer the candidate's domain of expertise based on their skills, professional summary, and work experience.

The domains can include:
- Technical domains: Data Engineering, Software Development, Web Development, etc.
- Non-technical domains: Finance and Accounting, Marketing, Sales, Human Resources, etc.
- Other domains: Healthcare, Education, Legal, etc.

Return the following fields in JSON format:

{
  "name": "<Candidate's Name>",
  "email": "<Candidate's Email>",
  "phone": "<Candidate's Phone Number>",
  "linkedin": "<Candidate's LinkedIn Profile URL>",
  "github": "<Candidate's GitHub Profile URL>",
  "professional_summary": "<Candidate's Professional Summary>",
  "domain": "<Inferred Domain (e.g., Data Engineering, Finance and Accounting)>",
  "technical_skills": ["<Skill 1>", "<Skill 2>", "..."],  # For technical roles
  "domain_skills": ["<Skill 1>", "<Skill 2>", "..."],    # For non-technical roles
  "work_experience": [
    {
      "role": "<Job Role>",
      "company": "<Company Name>",
      "start_date": "<Start Date>",
      "end_date": "<End Date>",
      "responsibilities": ["<Responsibility 1>", "<Responsibility 2>", "..."]
    },
    ...
  ],
  "education": [
    {
      "degree": "<Degree>",
      "institution": "<Institution Name>",
      "year": "<Year>"
    },
    ...
  ],
  "certifications": ["<Certification 1>", "<Certification 2>", "..."],
  "projects": [
    {
      "name": "<Project Name>",
      "description": "<Project Description>"
    },
    ...
  ]
}
Here is the candidate's resume:
"""

def fetch_candidate_details(candidate_id):
    conn = connect_to_postgres()
    cursor = conn.cursor()
    query = """
    SELECT name, email, phone, linkedin, github, professional_summary, domain
    FROM candidates
    WHERE candidate_id = %s;
    """
    cursor.execute(query, (candidate_id,))
    candidate_details = cursor.fetchone()
    conn.close()
    if candidate_details:
        keys = ["name", "email", "phone", "linkedin", "github", "professional_summary", "domain"]
        return dict(zip(keys, candidate_details))
    return None

def extract_job_description(file):
    try:
        if file.name.endswith(".docx"):
            doc = Document(file.name)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif file.name.endswith(".txt"):
            with open(file.name, "r") as f:
                return f.read()
        else:
            return "Unsupported file type. Excepted file types .txt or .docx file."
    except Exception as e:
        return f"Error reading file: {e}"
    

#PART 2 utility functions

def find_best_candidates(job_description):
    # Generate embedding for the job description
    job_embedding = embedding_model.encode(job_description)
    # print(job_embedding)
    # Query ChromaDB for the closest matches
    job_domain = extract_domain_from_description(job_description) 
    print(job_domain)
    print(f"The job is from {job_domain} domain")
    results = collection.query(
        query_embeddings=[job_embedding],
        n_results=3,  # Retrieve top 3 candidates
        # where={"domain": job_domain}
    )
    matched_candidates = []
    if results and results["documents"]:
        for meta, score in zip(results["metadatas"][0], results["distances"][0]):
            if "candidate_id" in meta and score <= 0.93:  # Filter by similarity score
                candidate_id = meta["candidate_id"]
                candidate_details = fetch_candidate_details(candidate_id)
                if candidate_details:
                    candidate_details["retrieval_score"] = score
                    matched_candidates.append(candidate_details)

    return matched_candidates

system_prompt2 = """
You are an assistant that predicts the domain the job description is about.

The domains can be:
- Technical domains: Data Engineering, Software Development, Web Development, etc.
- Non-technical domains: Finance and Accounting, Marketing, Sales, Human Resources, etc.
- Other domains: Healthcare, Education, Legal, etc.

Return just the name of the domain as a string:
"""

# Function to extract domain from job description using Ollama
def extract_domain_from_description(job_description):
    messages = [
        {"role": "system", "content": system_prompt2},
        {"role": "user", "content": job_description}
    ]
    response = ollama.chat(model=ollama_model, messages=messages)
    job_domain = response["message"]["content"].strip()
    return job_domain
